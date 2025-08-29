"""
README to Word Document Exporter

This script converts the README.md file to a professionally formatted Word document
with proper styling, headers, code blocks, and table of contents.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import markdown
from markdown.extensions import codehilite, tables, toc
from bs4 import BeautifulSoup


class ReadmeToWordExporter:
    """Export README.md to a professionally formatted Word document."""
    
    def __init__(self, readme_path: str, output_path: str = None):
        """
        Initialize the exporter.
        
        Args:
            readme_path (str): Path to the README.md file
            output_path (str): Output path for the Word document (optional)
        """
        self.readme_path = Path(readme_path)
        self.output_path = Path(output_path) if output_path else self.readme_path.parent / "README_Documentation.docx"
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Set up custom styles for the Word document."""
        styles = self.doc.styles
        
        # Helper function to add style safely
        def add_style_safe(name, style_type):
            try:
                return styles.add_style(name, style_type)
            except ValueError:
                return styles[name]
        
        # Title style
        title_style = add_style_safe('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        h1_style = add_style_safe('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        h2_style = add_style_safe('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Heading 3 style
        h3_style = add_style_safe('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Code block style
        code_style = add_style_safe('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Inline code style
        inline_code_style = add_style_safe('InlineCode', WD_STYLE_TYPE.CHARACTER)
        inline_code_style.font.name = 'Consolas'
        inline_code_style.font.size = Pt(10)
        
        # Quote style
        quote_style = add_style_safe('CustomQuote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.italic = True
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(6)
        quote_style.paragraph_format.space_after = Pt(6)
    
    def add_header_footer(self):
        """Add header and footer to the document."""
        # Header
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "User Authentication System - API Documentation"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Footer
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def process_markdown_line(self, line: str):
        """Process a single line of markdown and add to document."""
        line = line.rstrip()
        
        if not line:
            self.doc.add_paragraph()
            return
        
        # Headers
        if line.startswith('# '):
            self.doc.add_paragraph(line[2:], style='CustomTitle')
        elif line.startswith('## '):
            self.doc.add_paragraph(line[3:], style='CustomHeading1')
        elif line.startswith('### '):
            self.doc.add_paragraph(line[4:], style='CustomHeading2')
        elif line.startswith('#### '):
            self.doc.add_paragraph(line[5:], style='CustomHeading3')
        elif line.startswith('##### '):
            self.doc.add_paragraph(line[6:], style='CustomHeading3')
        
        # Code blocks
        elif line.startswith('```'):
            return 'code_block'
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            para = self.doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            para = self.doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        
        # Quotes
        elif line.startswith('> '):
            self.doc.add_paragraph(line[2:], style='CustomQuote')
        
        # Tables (simple detection)
        elif '|' in line and line.count('|') >= 2:
            return 'table_line'
        
        # Regular paragraph
        else:
            para = self.doc.add_paragraph(line)
            # Process inline formatting
            self.process_inline_formatting(para)
    
    def process_inline_formatting(self, paragraph):
        """Process inline formatting like bold, italic, and code."""
        text = paragraph.text
        
        # Process inline code first
        code_pattern = r'`([^`]+)`'
        matches = list(re.finditer(code_pattern, text))
        
        if matches:
            paragraph.clear()
            last_end = 0
            
            for match in matches:
                # Add text before code
                if match.start() > last_end:
                    paragraph.add_run(text[last_end:match.start()])
                
                # Add code text
                code_run = paragraph.add_run(match.group(1))
                code_run.style = self.doc.styles['InlineCode']
                
                last_end = match.end()
            
            # Add remaining text
            if last_end < len(text):
                paragraph.add_run(text[last_end:])
    
    def process_code_block(self, lines: list, start_index: int):
        """Process a code block and return the end index."""
        i = start_index + 1
        code_lines = []
        
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # Add code block to document
        if code_lines:
            code_text = '\n'.join(code_lines)
            para = self.doc.add_paragraph(code_text, style='CodeBlock')
            # Add shading to code blocks
            self.add_shading_to_paragraph(para, "F5F5F5")
        
        return i
    
    def add_shading_to_paragraph(self, paragraph, color: str):
        """Add background shading to a paragraph."""
        try:
            from docx.oxml import parse_xml
            shading_xml = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>'
            shading_elm = parse_xml(shading_xml)
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            # If shading fails, just continue without it
            pass
    
    def process_table(self, lines: list, start_index: int):
        """Process a markdown table and return the end index."""
        table_lines = []
        i = start_index
        
        # Collect table lines
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i])
            i += 1
        
        if len(table_lines) < 2:
            return start_index
        
        # Parse table
        header_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        data_rows = []
        
        for line in table_lines[2:]:  # Skip header separator
            if line.strip():
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                if row:
                    data_rows.append(row)
        
        # Create table in document
        if header_row and data_rows:
            table = self.doc.add_table(rows=1, cols=len(header_row))
            table.style = 'Table Grid'
            
            # Add header
            header_cells = table.rows[0].cells
            for j, header in enumerate(header_row):
                if j < len(header_cells):
                    header_cells[j].text = header
                    # Make header bold
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Add data rows
            for row_data in data_rows:
                row_cells = table.add_row().cells
                for j, cell_data in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = cell_data
        
        return i - 1
    
    def export(self) -> str:
        """
        Export the README.md to a Word document.
        
        Returns:
            str: Path to the generated Word document
        """
        if not self.readme_path.exists():
            raise FileNotFoundError(f"README file not found: {self.readme_path}")
        
        # Read the README file
        with open(self.readme_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        
        # Add header and footer
        self.add_header_footer()
        
        # Process the markdown content
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('```'):
                i = self.process_code_block(lines, i)
            elif '|' in line and line.count('|') >= 2:
                i = self.process_table(lines, i)
            else:
                self.process_markdown_line(line)
            
            i += 1
        
        # Add page break before any appendix
        self.doc.add_page_break()
        
        # Add document metadata
        self.add_document_info()
        
        # Save the document
        self.doc.save(self.output_path)
        
        return str(self.output_path)
    
    def add_document_info(self):
        """Add document information and metadata."""
        self.doc.add_heading('Document Information', level=1)
        
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ['Document Title', 'User Authentication System - API Documentation'],
            ['Generated Date', datetime.now().strftime('%B %d, %Y at %I:%M %p')],
            ['Source File', str(self.readme_path.name)],
            ['Document Type', 'Technical Documentation'],
            ['Version', '1.0']
        ]
        
        for i, (label, value) in enumerate(info_data):
            row_cells = info_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            # Make labels bold
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def main():
    """Main function to export README to Word document."""
    # Get the current directory
    current_dir = Path(__file__).parent
    readme_path = current_dir / "README.md"
    output_path = current_dir / "User_Authentication_System_Documentation.docx"
    
    try:
        # Create exporter and export
        exporter = ReadmeToWordExporter(readme_path, output_path)
        generated_file = exporter.export()
        
        print(f"✅ Successfully exported README to Word document!")
        print(f"📄 Generated file: {generated_file}")
        print(f"📁 File size: {os.path.getsize(generated_file) / 1024:.1f} KB")
        
        # Open the file location in Windows Explorer
        if os.name == 'nt':  # Windows
            os.startfile(str(Path(generated_file).parent))
        
    except FileNotFoundError as e:
        print(f"❌ Error: {e}")
    except Exception as e:
        print(f"❌ An error occurred: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
