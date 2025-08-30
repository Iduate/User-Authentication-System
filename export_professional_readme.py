"""
Professional README to Word Document Exporter

This script converts the comprehensive README.md file to a beautifully formatted 
Word document suitable for documentation, presentations, or project submissions.
"""

import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ProfessionalReadmeExporter:
    """Export README.md to a professional, comprehensive Word document."""
    
    def __init__(self, readme_path: str, output_path: str = None):
        """
        Initialize the exporter.
        
        Args:
            readme_path (str): Path to the README.md file
            output_path (str): Output path for the Word document (optional)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required. Install with: pip install python-docx")
        
        self.readme_path = Path(readme_path)
        if output_path:
            self.output_path = Path(output_path)
        else:
            # Create unique filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.output_path = self.readme_path.parent / f"User_Authentication_System_Complete_Documentation_{timestamp}.docx"
        
        self.doc = Document()
        self.setup_professional_styles()
    
    def setup_professional_styles(self):
        """Set up professional, corporate-style document formatting."""
        styles = self.doc.styles
        
        # Helper function to add style safely
        def add_style_safe(name, style_type):
            try:
                return styles.add_style(name, style_type)
            except ValueError:
                return styles[name]
        
        # Document title style
        title_style = add_style_safe('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(18)
        
        # Main heading style (H2)
        h1_style = add_style_safe('MainHeading', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(9)
        
        # Sub heading style (H3)
        h2_style = add_style_safe('SubHeading', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(68, 68, 68)  # Dark gray
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Minor heading style (H4)
        h3_style = add_style_safe('MinorHeading', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(68, 68, 68)  # Dark gray
        h3_style.paragraph_format.space_before = Pt(9)
        h3_style.paragraph_format.space_after = Pt(3)
        
        # Code block style
        code_style = add_style_safe('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.4)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Inline code style
        try:
            inline_code_style = add_style_safe('InlineCode', WD_STYLE_TYPE.CHARACTER)
            inline_code_style.font.name = 'Consolas'
            inline_code_style.font.size = Pt(10)
            inline_code_style.font.color.rgb = RGBColor(51, 51, 51)
        except:
            pass
        
        # Feature list style
        feature_style = add_style_safe('FeatureList', WD_STYLE_TYPE.PARAGRAPH)
        feature_style.font.name = 'Calibri'
        feature_style.font.size = Pt(11)
        feature_style.paragraph_format.left_indent = Inches(0.25)
        feature_style.paragraph_format.space_after = Pt(3)
    
    def add_professional_header_footer(self):
        """Add professional header and footer."""
        # Header
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "User Authentication System - Complete Technical Documentation"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.style.font.size = Pt(10)
        header_para.style.font.color.rgb = RGBColor(68, 68, 68)
        
        # Footer
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y')} | Django REST Framework Authentication Service"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.style.font.size = Pt(9)
        footer_para.style.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_table_of_contents(self):
        """Add a table of contents."""
        self.doc.add_heading('Table of Contents', level=1).style = 'MainHeading'
        
        toc_items = [
            "What it does",
            "Quick Start",
            "API Documentation", 
            "Local Development Setup",
            "Environment Variables",
            "Testing",
            "Security Features", 
            "Production Deployment",
            "Need Help?"
        ]
        
        for item in toc_items:
            para = self.doc.add_paragraph(f"• {item}")
            para.style.font.size = Pt(11)
            para.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
    
    def add_shading_to_paragraph(self, paragraph, color: str = "F8F9FA"):
        """Add light background shading to code blocks."""
        try:
            from docx.oxml import parse_xml
            shading_xml = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>'
            shading_elm = parse_xml(shading_xml)
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            pass
    
    def process_table(self, lines: list, start_index: int):
        """Process markdown tables into Word tables."""
        table_lines = []
        i = start_index
        
        # Collect table lines
        while i < len(lines) and ('|' in lines[i] or lines[i].strip() == ''):
            if lines[i].strip() and '|' in lines[i]:
                table_lines.append(lines[i])
            i += 1
        
        if len(table_lines) < 2:
            return start_index
        
        # Parse table header
        header_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1] if cell.strip()]
        
        # Parse table data (skip separator row)
        data_rows = []
        for line in table_lines[2:]:
            if line.strip():
                row = [cell.strip() for cell in line.split('|')[1:-1] if cell.strip()]
                if row and len(row) == len(header_row):
                    data_rows.append(row)
        
        # Create Word table
        if header_row and data_rows:
            table = self.doc.add_table(rows=1, cols=len(header_row))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Style header row
            header_cells = table.rows[0].cells
            for j, header in enumerate(header_row):
                if j < len(header_cells):
                    header_cells[j].text = header
                    for paragraph in header_cells[j].paragraphs:
                        paragraph.style.font.bold = True
                        paragraph.style.font.size = Pt(10)
            
            # Add data rows
            for row_data in data_rows:
                row_cells = table.add_row().cells
                for j, cell_data in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = cell_data
                        for paragraph in row_cells[j].paragraphs:
                            paragraph.style.font.size = Pt(10)
        
        return i - 1
    
    def process_markdown_content(self, content: str):
        """Process markdown content with enhanced formatting."""
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End of code block
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        para = self.doc.add_paragraph(code_text, style='CodeBlock')
                        self.add_shading_to_paragraph(para)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # Empty line
            if not line:
                # Only add paragraph if not following a heading
                if i > 0 and not lines[i-1].startswith('#'):
                    self.doc.add_paragraph()
                i += 1
                continue
            
            # Tables
            if '|' in line and line.count('|') >= 2:
                i = self.process_table(lines, i)
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                self.doc.add_paragraph(line[2:], style='DocumentTitle')
            elif line.startswith('## '):
                self.doc.add_paragraph(line[3:], style='MainHeading')
            elif line.startswith('### '):
                self.doc.add_paragraph(line[4:], style='SubHeading')
            elif line.startswith('#### '):
                self.doc.add_paragraph(line[5:], style='MinorHeading')
            
            # Feature bullet points with **bold** text
            elif line.startswith('- **') and line.endswith('**'):
                # Extract bold text
                bold_text = line[4:-2]  # Remove '- **' and '**'
                para = self.doc.add_paragraph()
                para.style = 'FeatureList'
                run = para.add_run(f"• {bold_text}")
                run.bold = True
            elif line.startswith('- **') and '**' in line:
                # Handle mixed bold and regular text
                para = self.doc.add_paragraph()
                para.style = 'FeatureList'
                para.add_run("• ")
                
                # Process bold parts
                parts = line[2:].split('**')
                for idx, part in enumerate(parts):
                    if idx % 2 == 1:  # Odd indices are bold
                        run = para.add_run(part)
                        run.bold = True
                    else:
                        para.add_run(part)
            
            # Regular bullet points
            elif line.startswith('- '):
                para = self.doc.add_paragraph(line[2:], style='List Bullet')
                self.process_inline_formatting(para)
            elif line.startswith('  - '):  # Sub-bullets
                para = self.doc.add_paragraph(line[4:], style='List Bullet 2')
                self.process_inline_formatting(para)
            
            # Numbered lists
            elif re.match(r'^\d+\. ', line):
                content = re.sub(r'^\d+\. ', '', line)
                para = self.doc.add_paragraph(content, style='List Number')
                self.process_inline_formatting(para)
            
            # Bold standalone lines (like **Using Docker:**)
            elif line.startswith('**') and line.endswith('**'):
                para = self.doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
            
            # Horizontal rule
            elif line.strip() == '---':
                para = self.doc.add_paragraph('_' * 70)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Regular paragraph
            else:
                para = self.doc.add_paragraph(line)
                self.process_inline_formatting(para)
            
            i += 1
    
    def process_inline_formatting(self, paragraph):
        """Process inline formatting like code, bold, and links."""
        text = paragraph.text
        
        # Process inline code (`code`)
        if '`' in text:
            code_pattern = r'`([^`]+)`'
            matches = list(re.finditer(code_pattern, text))
            
            if matches:
                paragraph.clear()
                last_end = 0
                
                for match in matches:
                    # Add text before code
                    if match.start() > last_end:
                        paragraph.add_run(text[last_end:match.start()])
                    
                    # Add code text
                    code_run = paragraph.add_run(match.group(1))
                    try:
                        code_run.style = self.doc.styles['InlineCode']
                    except:
                        code_run.font.name = 'Consolas'
                        code_run.font.size = Pt(10)
                    
                    last_end = match.end()
                
                # Add remaining text
                if last_end < len(text):
                    paragraph.add_run(text[last_end:])
    
    def add_cover_page(self):
        """Add a professional cover page."""
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(72)
        
        title_run = title_para.add_run("User Authentication System")
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle_para = self.doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.paragraph_format.space_before = Pt(12)
        
        subtitle_run = subtitle_para.add_run("Complete Technical Documentation")
        subtitle_run.font.name = 'Calibri'
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = RGBColor(68, 68, 68)
        
        # Description
        desc_para = self.doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_para.paragraph_format.space_before = Pt(36)
        
        desc_run = desc_para.add_run("Django REST Framework Authentication Service\nwith JWT Tokens, Rate Limiting, and Docker Support")
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(14)
        desc_run.font.italic = True
        
        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Pt(72)
        
        date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(12)
        
        self.doc.add_page_break()
    
    def export(self) -> str:
        """
        Export the README.md to a professional Word document.
        
        Returns:
            str: Path to the generated Word document
        """
        if not self.readme_path.exists():
            raise FileNotFoundError(f"README file not found: {self.readme_path}")
        
        # Read the README file
        with open(self.readme_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Add cover page
        self.add_cover_page()
        
        # Add table of contents
        self.add_table_of_contents()
        
        # Add header and footer
        self.add_professional_header_footer()
        
        # Process the markdown content
        self.process_markdown_content(content)
        
        # Add document metadata page
        self.doc.add_page_break()
        self.add_document_metadata()
        
        # Save the document
        self.doc.save(self.output_path)
        
        return str(self.output_path)
    
    def add_document_metadata(self):
        """Add comprehensive document metadata."""
        self.doc.add_heading('Document Information', level=1).style = 'MainHeading'
        
        # Create metadata table
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        metadata = [
            ['Document Title', 'User Authentication System - Complete Technical Documentation'],
            ['Project Type', 'Django REST Framework Authentication Service'],
            ['Generated Date', datetime.now().strftime('%B %d, %Y at %I:%M %p')],
            ['Source File', str(self.readme_path.name)],
            ['Technologies', 'Django, PostgreSQL, Redis, Docker, JWT'],
            ['API Documentation', 'Swagger/OpenAPI available'],
            ['Test Coverage', 'Unit tests for all authentication flows'],
            ['Deployment Ready', 'Railway, Render, Heroku compatible']
        ]
        
        for i, (label, value) in enumerate(metadata):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            
            # Style the cells
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            
            for paragraph in row_cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def main():
    """Main function to export README to professional Word document."""
    current_dir = Path(__file__).parent
    readme_path = current_dir / "README.md"
    
    try:
        if not DOCX_AVAILABLE:
            print("❌ Missing required package!")
            print("📦 Please install python-docx first:")
            print("   pip install python-docx")
            return
        
        # Create exporter and export
        exporter = ProfessionalReadmeExporter(readme_path)
        generated_file = exporter.export()
        
        print("✅ Successfully exported README to professional Word document!")
        print(f"📄 Generated file: {generated_file}")
        print(f"📁 File size: {os.path.getsize(generated_file) / 1024:.1f} KB")
        print("\n📋 Document includes:")
        print("   • Professional cover page")
        print("   • Table of contents")
        print("   • Complete API documentation")
        print("   • Setup instructions")
        print("   • Environment variables reference")
        print("   • Testing and deployment guides")
        print("   • Professional formatting and styling")
        
        # Try to open the file location
        if os.name == 'nt':  # Windows
            try:
                os.startfile(str(Path(generated_file).parent))
                print("📂 Opened file location in Explorer")
            except:
                print(f"📂 File saved to: {Path(generated_file).parent}")
        
    except FileNotFoundError as e:
        print(f"❌ Error: {e}")
    except ImportError as e:
        print(f"❌ Import Error: {e}")
    except Exception as e:
        print(f"❌ An error occurred: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
